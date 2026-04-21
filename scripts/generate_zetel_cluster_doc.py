from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

OUTPUT_PATH = Path("/Users/filipbreckx/Development/interieur.expert/Zetel-cluster-dossier-en-artikels.docx")

DOC_TITLE = "Zetelcluster Vlaanderen - dossierpagina en longform artikels"

INTRO = """
# Zetelcluster Vlaanderen

**Doel van dit document:** een volledig uitgewerkt contentpakket voor een Vlaams SEO- en GEO-cluster rond zetels kopen. Dit document bevat de dossierpagina, alle longform ondersteunende artikels en een uitgewerkte videopagina voor de bestaande YouTube-video met koopadvies.

**Strategische insteek:** focus op mensen in Vlaanderen die een nieuwe zetel willen kopen en vooral advies zoeken, niet louter prijsvergelijking. Hoofdterm is **zetel**, met **sofa** en **bank** als natuurlijke semantische varianten in tekst, tussenkoppen en ankerlinks.

**Gebruikte SEO-opportuniteiten:** waar moet je op letten, welke zetel past bij mij, stof of leder, kwaliteit herkennen, afmetingen, testen in de winkel, showroom- en koopadvies. Dit sluit aan bij terugkerende SERP-patronen op Belgische en Nederlandse marktspelers, maar is inhoudelijk vertaald naar een neutralere, redactionele en Vlaams georiënteerde invalshoek.

## Aanbevolen publicatievolgorde

1. **Dossierpagina:** /dossiers/zetel-kiezen
2. **Pillar:** /advies/zetel-kopen-waarop-letten
3. **Videopagina:** /video/tips-voor-het-kopen-van-een-zetel
4. **Ondersteunend artikel:** /advies/zetel-testen-in-de-winkel
5. **Ondersteunend artikel:** /advies/stof-of-leren-zetel-kiezen
6. **Ondersteunend artikel:** /advies/welke-zetel-past-bij-jou
7. **Ondersteunend artikel:** /advies/kwalitatieve-zetel-herkennen
8. **Ondersteunend artikel:** /advies/zetel-afmetingen-woonkamer

> Wie vandaag zoekt op termen als "zetel kopen tips" of "waarop letten bij zetel kopen", wil vooral fouten vermijden. Dat maakt praktische, concrete en geloofwaardige inhoud belangrijker dan algemene lifestyle-copy.
""".strip()

SECTIONS: list[tuple[str, str]] = [
    (
        "Dossierpagina",
        """
# Dossierpagina

**Titel:** Zetel kiezen: complete gids van comfort tot kwaliteit  
**Slug:** /dossiers/zetel-kiezen  
**Primaire zoekintentie:** oriëntatie + commerciële informatie  
**Primaire keywordfocus:** zetel kiezen, zetel kopen, waar op letten bij zetel kopen  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Welke zetel past bij jou?](https://interieur.expert/advies/welke-zetel-past-bij-jou), [Stof of lederen zetel?](https://interieur.expert/advies/stof-of-leren-zetel-kiezen), [Hoe herken je een kwalitatieve zetel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen), [Welke afmeting van zetel werkt in jouw woonkamer?](https://interieur.expert/advies/zetel-afmetingen-woonkamer), [Een zetel testen in de winkel](https://interieur.expert/advies/zetel-testen-in-de-winkel), [Modulaire sofa kiezen](https://interieur.expert/advies/modulaire-sofa-kiezen-voordelen-tips)  
**Externe links:** [Livios - inspiratie rond wonen](https://www.livios.be/nl/), [Test Aankoop](https://www.test-aankoop.be/)

**Excerpt:** Van comfort en afmetingen tot stof, leder en kwaliteit: in dit dossier ontdek je stap voor stap hoe je een zetel kiest die goed zit, goed oogt en lang meegaat.

Een zetel kies je niet voor een seizoen. Je leeft ermee, zit erin, ontvangt er bezoek op en kijkt er elke dag naar. Toch wordt die keuze vaak te snel gemaakt. Mooie stof, goed gevoel in de winkel, klaar. Alleen: wat in een showroom goed oogt, past niet altijd bij hoe je thuis echt woont.

In dit dossier brengen we alle belangrijke keuzes samen. Niet vanuit trends alleen, maar vanuit gebruik, comfort en lange termijn. Want een goede zetel is geen impulsmeubel. Het is een stuk dat je leven makkelijker en aangenamer moet maken.

> De juiste zetel herken je niet aan één detail, maar aan het totaal: hoe hij zit, hoe hij in je ruimte past en hoe geloofwaardig hij gebouwd is.

## Waar draait de keuze echt om?

Wie online zoekt naar een nieuwe zetel, zoekt zelden alleen naar kleur of vorm. Meestal zit er een concretere vraag achter. Past een hoekzetel wel in mijn woonkamer? Is stof wel slim met kinderen? Hoe weet ik of een zetel degelijk is? Moet ik stevig of zachter zitten kiezen? Precies daarom is een goede keuze zelden puur esthetisch.

Een zetel moet eerst kloppen met je dagelijks leven. Zit je elke avond languit te lezen of tv te kijken, dan heb je andere noden dan iemand die vooral een formele zithoek wil. Ontvang je vaak bezoek, dan tellen circulatie en opstelling zwaarder mee. Heb je jonge kinderen of een hond, dan wordt onderhoud belangrijker dan bij een rustige leefruimte voor twee.

## Zo gebruik je dit dossier

Werk niet van kleur naar comfort, maar omgekeerd. Start bij hoe je woont, dan bij formaat, daarna bij materiaal en pas op het einde bij uitstraling. Dat voorkomt dat je verliefd wordt op een model dat in de praktijk niet werkt.

In de artikels hieronder helpen we je stap voor stap:

- Eerst krijg je de brede koopgids met alles waar je op moet letten.
- Daarna zoomen we in op levensstijl, materiaalkeuze en kwaliteit.
- Vervolgens kijken we naar afmetingen en wat je in de winkel echt moet testen.
- De videopagina geeft extra geloofwaardigheid, omdat ze koopadvies bundelt vanuit de praktijk van een meubelverkoper.

## Welke fouten maken kopers het vaakst?

De grootste fout is kiezen op beeld alleen. Een zetel kan op foto perfect ogen en in de winkel degelijk lijken, maar thuis te diep, te log of te onderhoudsgevoelig blijken. De tweede fout is onderschatten hoeveel plaats een zetel echt vraagt. Niet alleen de meubelmaat telt, maar ook de loopruimte eromheen. De derde fout is dat mensen kwaliteit proberen af te lezen aan prijs of merknaam alleen, terwijl opbouw, vering en afwerking vaak veel meer zeggen.

Ook comfort wordt vaak te snel beoordeeld. Twee minuten zitten is niet genoeg. Een zetel moet goed voelen in verschillende houdingen: rechtop, onderuit, met steun in de rug en bij het opstaan. Zeker voor een aankoop die jaren meegaat, loont het om bewuster te testen.

## Wat maakt dit cluster sterk voor SEO en GEO?

Dit dossier is opgezet als centrale ingang voor alles rond zetels kopen. Het vangt de brede oriëntatievraag op en leidt bezoekers logisch verder naar diepere deelvragen. Dat is niet alleen nuttig voor lezers, maar ook voor zoekmachines en AI-systemen die proberen te begrijpen welk stuk inhoud de centrale bron is binnen een onderwerp.

Door alle ondersteunende artikels rond duidelijke keuzevragen te organiseren, ontstaat een cluster waarin elk stuk een eigen zoekintentie bedient, terwijl het dossier overzicht en samenhang brengt. Dat verhoogt de kans dat meerdere pagina's tegelijk zichtbaar worden op gerelateerde queries.

## De artikels in dit dossier

### 1. Zetel kopen: waar moet je echt op letten?
De brede gids voor iedereen die aan de start van het keuzeproces staat.

### 2. Welke zetel past bij jouw leven?
Voor wie nog twijfelt tussen modellen, gebruikssituaties en woonprofielen.

### 3. Stof of lederen zetel?
Een klassieke maar belangrijke keuze, vaak met hoge koopintentie.

### 4. Hoe herken je een kwalitatieve zetel?
Een trust-artikel dat helpt om showroompraat te onderscheiden van echte kwaliteit.

### 5. Welke afmeting van zetel werkt in jouw woonkamer?
Praktische evergreen met sterke long-tail potentie.

### 6. Een zetel testen in de winkel
Perfect koppelbaar aan video en heel bruikbaar voor bezoekers vlak voor aankoop.

### 7. Video: tips voor het kopen van een zetel
Een ondersteunende asset met menselijke geloofwaardigheid en sterke dwell-time waarde.

## Slot

Wie een zetel goed kiest, koopt niet gewoon een meubel. Je kiest hoe je ontspant, hoe je woonkamer aanvoelt en hoe comfortabel je dagelijks leeft. Daarom loont het om de aankoop trager, slimmer en inhoudelijker aan te pakken dan veel showroombezoeken of snelle online vergelijkingen suggereren.

Gebruik dit dossier als vertrekpunt. Werk van gebruik naar formaat, van formaat naar materiaal en van materiaal naar kwaliteit. Dan wordt de kans veel groter dat je niet alleen een mooie zetel koopt, maar vooral een juiste.
""".strip(),
    ),
    (
        "Pillarartikel",
        """
# Artikel

**Titel:** Zetel kopen: waar moet je echt op letten?  
**Slug:** /advies/zetel-kopen-waarop-letten  
**Primaire zoekintentie:** informatief met sterke commerciële intentie  
**Primaire keywordfocus:** zetel kopen, waarop letten bij zetel kopen, tips zetel kopen  
**Interne links:** [Dossier: Zetel kiezen](https://interieur.expert/dossiers/zetel-kiezen), [Welke zetel past bij jou?](https://interieur.expert/advies/welke-zetel-past-bij-jou), [Stof of lederen zetel?](https://interieur.expert/advies/stof-of-leren-zetel-kiezen), [Hoe herken je een kwalitatieve zetel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen), [Een zetel testen in de winkel](https://interieur.expert/advies/zetel-testen-in-de-winkel), [Video: tips voor het kopen van een zetel](https://interieur.expert/video/tips-voor-het-kopen-van-een-zetel)

**Excerpt:** Een zetel koop je voor jaren. Ontdek waar je echt op moet letten qua comfort, formaat, materiaal, kwaliteit en showroomtests zodat je geen dure miskoop doet.

Een nieuwe zetel kopen lijkt eenvoudig tot je er echt aan begint. In de winkel ziet bijna elk model er comfortabel uit. Online lijken stoffen, afmetingen en kleuren vaak duidelijker dan ze in werkelijkheid zijn. En dan is er nog de klassieke twijfel: ga je voor een hoekzetel, een compacte 3-zit, een modulaire sofa of toch een strakker model dat minder plaats inneemt?

Het probleem is zelden een gebrek aan keuze. Integendeel. Net doordat er zoveel modellen, materialen en comforttypes bestaan, kiezen veel mensen te snel op uitstraling. Daardoor kopen ze een zetel die mooi oogt, maar niet goed past bij hun ruimte, hun gewoontes of hun verwachtingen op lange termijn.

> Een goede zetel koop je niet op basis van één mooie showroomopstelling, maar op basis van hoe je thuis echt leeft.

## Begin niet bij kleur, maar bij gebruik

De eerste vraag is niet welke zetel je mooi vindt, maar hoe je hem gebruikt. Zit je elke avond languit onder een dekentje? Ontvang je vaak vrienden? Heb je kinderen die over de zetel kruipen? Werk je soms met een laptop in de woonkamer? Elk van die situaties vraagt iets anders van zitdiepte, opstelling, bekleding en onderhoud.

Wie vooral wil loungen, zoekt meestal meer diepte en zachtere ondersteuning. Wie rechtop zit om te lezen of bezoek te ontvangen, heeft eerder baat bij een actievere zit en voldoende steun in de rug. Een zetel voor een rustige designleefruimte vraagt dus niet hetzelfde als een gezinszetel die elke dag intensief gebruikt wordt.

Daarom loont het om je levensstijl eerst scherp te krijgen. Pas daarna wordt het zinvol om modellen te vergelijken.

## De afmetingen bepalen meer dan je denkt

Veel miskopen ontstaan niet door slechte smaak, maar door verkeerde proporties. Een zetel kan in de winkel perfect aanvoelen en thuis plots veel te zwaar ogen. Dat gebeurt vooral wanneer mensen alleen naar de breedte kijken. Minstens even belangrijk zijn diepte, hoogte, armleuningen en de ruimte rond het meubel.

Een diepe loungezetel voelt uitnodigend, maar neemt visueel en fysiek meer ruimte in. Een model met brede armleuningen oogt luxueus, maar vreet zitruimte. Een lage zetel kan stijlvol zijn, maar werkt niet voor iedereen even comfortabel. Zeker in compactere woonkamers moet je niet alleen het meubel meten, maar ook de loopruimte naar deuren, ramen en andere zitplaatsen meenemen.

Wie daar dieper op wil ingaan, leest best ook [Welke afmeting van zetel werkt in jouw woonkamer?](https://interieur.expert/advies/zetel-afmetingen-woonkamer).

## Comfort is meer dan zacht zitten

Veel mensen verwarren comfort met zachtheid. Toch is een erg zachte zetel niet automatisch de beste keuze. Wat eerst aangenaam voelt, kan na een halfuur te weinig steun geven. Omgekeerd kan een iets stevigere zetel in dagelijks gebruik net beter zitten en mooier blijven.

Let daarom op drie dingen: hoe je zit, hoe je ondersteund wordt en hoe makkelijk je weer opstaat. Een goede rugondersteuning voel je niet als druk, maar als natuurlijke steun. De zitdiepte moet toelaten dat je ontspannen zit zonder dat je onderuit schuift of een kussen in je onderrug nodig hebt. En wie niet piepjong is of vaak actief opstaat, merkt snel of een zetel te laag of te zacht is.

## Stof, leder of iets ertussen?

Materiaal bepaalt niet alleen de uitstraling, maar ook hoe een zetel leeft in jouw huis. Stof voelt meestal warmer en huiselijker aan. Leder oogt vaak strakker en luxueuzer, maar vraagt een ander soort onderhoud. Ook performancestoffen, microvezels en moderne bekledingen verdienen aandacht, zeker bij intensief gebruik.

De juiste keuze hangt af van temperatuur, gezinssituatie, huisdieren, onderhoudsbereidheid en stijl. Daarom behandelen we die afweging apart in [Stof of lederen zetel? Zo maak je de juiste keuze](https://interieur.expert/advies/stof-of-leren-zetel-kiezen).

## Kwaliteit zie je niet alleen aan de buitenkant

Een zetel koop je vaak voor jaren, soms voor een decennium of langer. Dan volstaat het niet dat de bekleding mooi is of dat de showroomprijs vertrouwen uitstraalt. Kijk ook naar de opbouw: het frame, de vering, de vulling, de afwerking van naden en het herstel van de kussens na belasting.

Vraag gerust hoe het frame opgebouwd is, welke vering gebruikt wordt en hoe de zitkussens zijn samengesteld. Een degelijke verkoper kan daar duidelijk op antwoorden. Wie enkel over stofkleuren en promoties praat, helpt je niet noodzakelijk bij de juiste keuze.

Lees ook [Hoe herken je een kwalitatieve zetel in de winkel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen) als je showroompraat van echte kwaliteit wilt onderscheiden.

## Test in de winkel zoals je thuis leeft

Ga niet alleen even zitten en weer rechtstaan. Test een zetel in verschillende houdingen. Zit rechtop. Leun achterover. Schuif wat onderuit als je dat thuis ook doet. Leg je arm op de leuning. Kijk of je voeten natuurlijk op de grond staan. Sta meerdere keren op. Laat ook je partner of andere gebruiker testen, want comfort is persoonlijk.

Nog beter: neem de tijd en vergelijk twee of drie modellen kort na elkaar. Dan voel je sneller wat het verschil is tussen aangenaam eerste gevoel en echt duurzaam comfort. Voor een praktische checklist kun je later doorlinken naar [Een zetel testen in de winkel: 9 dingen die je echt moet doen](https://interieur.expert/advies/zetel-testen-in-de-winkel).

> De beste showroomtest is niet de meest elegante, maar de eerlijkste. Doe in de winkel wat je thuis ook zou doen.

## Kies een model dat past bij je ruimte, niet alleen bij je smaak

Veel mensen worden verliefd op één model zonder rekening te houden met de rest van de ruimte. Toch moet een zetel samenwerken met vloer, salontafel, lichtinval, zichtlijnen en circulatie. Een royale hoekzetel kan heerlijk zitten, maar een kamer zwaar maken. Een compact model kan ruimtelijk slim zijn, maar te weinig comfort bieden voor hoe je werkelijk woont.

Net daarom is evenwicht belangrijker dan extremen. Een zetel moet voldoende aanwezig zijn om de woonkamer te dragen, maar niet zo dominant dat alles errond verdwijnt. Vaak is de beste keuze niet de grootste of meest opvallende, maar de zetel die het meest vanzelfsprekend in de ruimte landt.

## Denk verder dan vandaag

Je koopt een zetel meestal niet alleen voor je huidige smaak, maar ook voor de komende jaren. Denk dus aan veranderende gewoontes, slijtage, onderhoud en mogelijke verhuis. Een heel trendgevoelig model kan na korte tijd vermoeien. Een neutrale basis met karaktervolle textuur of goed gekozen kussens blijft vaak langer relevant.

Dat betekent niet dat je veilig of saai moet kiezen. Wel dat kwaliteit, proportie en gebruik zwaarder mogen doorwegen dan het wow-effect van het moment.

## Conclusie

Wie een zetel goed wil kiezen, moet verder kijken dan vorm en stof. De juiste zetel past bij je leven, je ruimte en je verwachtingen op lange termijn. Start bij gebruik, test comfort bewust, kijk kritisch naar materiaal en opbouw, en durf vragen stellen in de winkel.

Dat vraagt misschien iets meer tijd, maar het verschil tussen een snelle aankoop en een goede keuze voel je jarenlang. En precies daarom is een zetel een aankoop die inhoudelijk advies verdient.
""".strip(),
    ),
    (
        "Levensstijlartikel",
        """
# Artikel

**Titel:** Welke zetel past bij jouw leven? Kies op basis van gezin, ruimte en gebruik  
**Slug:** /advies/welke-zetel-past-bij-jou  
**Primaire zoekintentie:** vergelijkend en oriënterend  
**Primaire keywordfocus:** welke zetel past bij mij, welke zetel kiezen, beste zetel voor gezin  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Zetel afmetingen woonkamer](https://interieur.expert/advies/zetel-afmetingen-woonkamer), [Stof of lederen zetel?](https://interieur.expert/advies/stof-of-leren-zetel-kiezen)

**Excerpt:** De beste zetel hangt af van hoe je woont. Ontdek welk type zetel logisch is voor gezinnen, compacte woonkamers, loungers, designliefhebbers en intensief dagelijks gebruik.

Er bestaat niet zoiets als de beste zetel in absolute zin. Er bestaat wel een zetel die het best past bij jouw manier van wonen. Dat klinkt vanzelfsprekend, maar toch vertrekken veel kopers vanuit model, trend of kleur in plaats van vanuit gebruik. Daardoor kiezen ze een zetel die misschien mooi oogt, maar in de praktijk niet logisch voelt.

Een gezin met jonge kinderen heeft andere noden dan een koppel dat vooral 's avonds rustig leest. Iemand met een compacte stadswoonkamer zoekt andere proporties dan iemand met een open leefruimte. En wie een zetel vooral gebruikt om languit tv te kijken, heeft weinig aan een formeel model dat vooral goed oogt op foto.

## Als je een gezin met kinderen hebt

In een gezinssituatie moet een zetel vooral veel aankunnen. Dan telt niet alleen comfort, maar ook onderhoud, opstelling en vergevingsgezindheid van het materiaal. Te lichte of erg gevoelige stoffen kunnen snel stress geven. Een model met losse kussens vraagt meer opschudden. Een heel lage designzetel kan moeilijker zijn bij dagelijks intens gebruik.

Vaak werken middelhoge modellen met een degelijke performancestof of onderhoudsvriendelijke bekleding hier het best. Ook een opstelling die voldoende zitplaatsen biedt zonder de kamer te blokkeren, is cruciaal. Denk dus niet alleen aan hoeveel mensen er kunnen zitten, maar ook aan hoe kinderen langs de zetel bewegen en erop spelen.

## Als je compact woont

In een kleinere woonkamer is de reflex vaak om zo klein mogelijk te kiezen. Dat is niet altijd slim. Een te kleine zetel kan de ruimte ondermaats en onaf maken. Veel belangrijker dan pure breedte zijn lucht, poothoogte, armleuningen en visuele massa.

Een model op fijne poten oogt vaak lichter dan een lage blokvorm. Smalle armleuningen geven meer bruikbare zitbreedte. En een iets compactere 3-zit kan soms logischer zijn dan een hoekzetel die de doorstroom belemmert. In kleine ruimtes is proportie belangrijker dan bravoure.

## Als je vooral wilt loungen

Wie languit wil liggen, series kijkt of graag onderuit zit, zal snel uitkomen bij diepere en zachtere modellen. Dat is logisch, maar let op dat loungecomfort niet altijd hetzelfde is als universeel comfort. Een diepe zetel die voor jou heerlijk is, kan voor bezoek of oudere gezinsleden minder praktisch zijn.

Daarom is het goed om te testen of de zetel ook zonder extra kussen steun geeft in de rug en of je gemakkelijk weer recht geraakt. Een loungezetel mag zacht zijn, maar hoeft niet vormloos te worden.

> Het beste comfort is niet per se het zachtste comfort, maar het comfort dat klopt voor de manier waarop je echt zit.

## Als uitstraling en design zwaar doorwegen

Sommige kopers willen een zetel die de ruimte esthetisch draagt. Dat is geen ijdelheid, maar een legitieme ontwerpkeuze. Alleen moet design ook hier samenwerken met gebruik. Een sculpturale zetel met prachtige lijnen kan perfect zijn in een rustig interieur, zolang hij niet te weinig steun geeft of te gevoelig is voor het dagelijkse leven dat errond gebeurt.

Voor designgerichte kopers werkt het vaak goed om neutraal te kiezen in kleur en gedurfder in vorm, of omgekeerd. Zo blijft het model uitgesproken zonder snel te vermoeien.

## Als je vaak bezoek ontvangt

Dan speelt opstelling een grotere rol. Een zetel moet niet alleen goed zitten, maar ook het gesprek ondersteunen. Heel diepe loungemodellen zijn heerlijk voor twee, maar minder geschikt als je vaak met meerdere mensen samen zit. Een actievere zit met voldoende rugondersteuning werkt dan vaak beter.

Ook de verhouding tot fauteuils, bijzettafels en de rest van de zithoek telt mee. Een goede bezoekzetel is zelden te laag, te diep of te gesloten in vorm.

## Als duurzaamheid en lange levensduur belangrijk zijn

Dan verschuift de focus naar tijdloze vormen, onderhoudsvriendelijke materialen en geloofwaardige opbouw. Wie bewust koopt, kiest vaak beter voor een model dat over vijf jaar nog goed zit en goed oogt, dan voor een trendstuk dat snel veroudert.

Dat betekent meestal: neutrale basis, degelijke stoffering, goed binnenwerk en een vorm die niet alleen vandaag aanspreekt.

## Welke conclusie trek je hieruit?

De juiste zetel begint bij een eerlijke blik op je leven. Woon je met kinderen, compact, formeel, relaxed of uitgesproken designgericht? Zodra dat helder is, vallen veel modellen vanzelf af en wordt de keuze veel logischer.

Dat maakt zetel kopen niet alleen makkelijker, maar ook slimmer. Je koopt dan geen zetel voor het beeld in de showroom, maar voor het leven dat zich thuis afspeelt.
""".strip(),
    ),
    (
        "Materiaalartikel",
        """
# Artikel

**Titel:** Stof of lederen zetel? Zo maak je de juiste keuze  
**Slug:** /advies/stof-of-leren-zetel-kiezen  
**Primaire zoekintentie:** materiaalvergelijking  
**Primaire keywordfocus:** stof of lederen zetel, stoffen zetel of lederen zetel, beste bekleding zetel  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Hoe herken je een kwalitatieve zetel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen), [Welke zetel past bij jouw leven?](https://interieur.expert/advies/welke-zetel-past-bij-jou)

**Excerpt:** Twijfel je tussen stof en leder voor je nieuwe zetel? Dit artikel vergelijkt comfort, onderhoud, uitstraling, slijtage en gebruik zodat je een keuze maakt die thuis ook echt werkt.

De keuze tussen stof en leder lijkt op het eerste gezicht eenvoudig: stof voelt warmer, leder oogt luxueuzer. Maar wie een zetel koopt voor dagelijks gebruik, weet dat het verschil dieper gaat. Materiaal bepaalt hoe een zetel aanvoelt, hoe hij veroudert, hoeveel onderhoud hij vraagt en hoe ontspannen je ermee leeft.

Er is dus geen universeel beter antwoord. Wel is er een materiaal dat logischer is voor jouw woonstijl, gezinssituatie en comfortverwachting.

## Waarom stof zo populair blijft

Een stoffen zetel voelt meestal meteen toegankelijk aan. Stof oogt huiselijk, warm en zacht. Dat maakt het voor veel woonkamers een vanzelfsprekende keuze. Bovendien is de variatie enorm: van strak geweven en bijna architecturaal tot losser, tactiel en gezellig.

Voor gezinnen en intensief dagelijks gebruik kan stof heel sterk zijn, zeker wanneer je kiest voor moderne performancestoffen of onderhoudsvriendelijke kwaliteiten. Ze verbergen kleine gebruikssporen vaak beter dan leder en voelen minder koud aan wanneer je gaat zitten.

## Waar je bij stof op moet letten

Niet elke stof is even praktisch. Een heel open geweven of erg delicate bekleding kan sneller gebruikssporen tonen. Ook kleur speelt mee. Heel lichte tinten ogen fris, maar zijn minder vergevingsgezind in een druk huishouden. Vraag daarom niet alleen welke stof mooi is, maar ook hoe ze zich gedraagt bij dagelijks gebruik.

Belangrijk zijn slijtvastheid, structuur, onderhoudbaarheid en hoe makkelijk de stof gereinigd kan worden. Zeker bij kinderen of huisdieren loont het om die praktische laag serieus te nemen.

## Waarom leder voor veel kopers aantrekkelijk blijft

Leder heeft een uitstraling die weinig andere materialen evenaren. Het oogt helder, krachtig en vaak ook rustiger in een interieur. Zeker in strakkere, warm-minimalistische of meer architecturale woonkamers kan een lederen zetel heel goed landen.

Daarnaast ontwikkelt echt leder karakter. Het leeft mee met gebruik en krijgt na verloop van tijd een patina dat door veel mensen net gewaardeerd wordt. Daardoor voelt een goede lederen zetel vaak als een stuk dat mooier wordt met de jaren.

> Wie voor leder kiest, kiest niet alleen voor een look, maar ook voor een materiaal dat zichtbaar ouder wordt. De vraag is dus niet of het verandert, maar of je die evolutie mooi vindt.

## Waar je bij leder kritisch op moet blijven

Leder is niet automatisch praktischer. Het vraagt een ander soort onderhoud en reageert op warmte, zonlicht en uitdroging. Ook het gevoel verschilt: in de winter kan het koeler aanvoelen, in de zomer warmer. Sommige mensen vinden dat storend, anderen net niet.

Bovendien is niet elk lederen uiterlijk gelijk. Afwerking, nerf, soepelheid en beschermlaag bepalen hoe natuurlijk of juist hoe gecorrigeerd het materiaal oogt. Vraag dus altijd welk type leder je koopt en hoe het onderhouden moet worden.

## Wat is slimmer met kinderen of huisdieren?

Dat hangt af van je tolerantie en je huishouden. Een onderhoudsvriendelijke stof kan hier perfect werken, zeker wanneer ze robuust en goed reinigbaar is. Leder kan vlekbestendig lijken, maar is niet immuun voor krassen, uitdroging of sporen van intensief gebruik.

De beste keuze is vaak niet het materiaal met de sterkste reputatie, maar het materiaal dat het best past bij hoe ontspannen je ermee wilt leven. Als je bij elk spoor onrust voelt, is zelfs een technisch degelijk materiaal geen goede keuze.

## Wat oogt tijdlozer?

Zowel stof als leder kunnen tijdloos zijn. De echte tijdloosheid zit minder in het materiaal zelf en meer in kleur, model en proportie. Een rustige stoffen zetel in een warme neutrale tint kan veel langer meegaan in uitstraling dan een modieuze lederen uitvoering. Omgekeerd kan een goed ontworpen lederen zetel bijzonder duurzaam en stijlvol blijven.

## Hoe kies je dan wel?

Stel jezelf drie vragen. Hoe wil ik dat de zetel aanvoelt? Hoe wil ik ermee leven? En hoe mag hij ouder worden? Wie een warm, zacht en gezinsvriendelijk gevoel zoekt, komt vaak bij stof uit. Wie rust, helderheid en een sterker patina waardeert, voelt zich sneller goed bij leder.

De juiste keuze is degene die thuis logisch blijft aanvoelen, ook na het eerste showroommoment.
""".strip(),
    ),
    (
        "Kwaliteitsartikel",
        """
# Artikel

**Titel:** Hoe herken je een kwalitatieve zetel in de winkel?  
**Slug:** /advies/kwalitatieve-zetel-herkennen  
**Primaire zoekintentie:** kwaliteit beoordelen voor aankoop  
**Primaire keywordfocus:** kwalitatieve zetel herkennen, goede kwaliteit zetel, waar op letten kwaliteit zetel  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Een zetel testen in de winkel](https://interieur.expert/advies/zetel-testen-in-de-winkel), [Stof of lederen zetel?](https://interieur.expert/advies/stof-of-leren-zetel-kiezen)

**Excerpt:** Een zetel kan in de showroom perfect ogen en toch middelmatig gebouwd zijn. Zo herken je degelijke kwaliteit aan frame, vering, vulling, afwerking en de juiste vragen aan de verkoper.

Een zetel beoordelen op kwaliteit is moeilijker dan veel mensen denken. Net omdat de buitenkant vaak overtuigend is, lijken veel modellen beter dan ze werkelijk zijn. Mooie stof, zachte kussens en goede styling verbergen makkelijk wat eronder zit. Toch zit de echte kwaliteit net in dat onzichtbare deel.

Wie slim koopt, leert dus niet alleen kijken naar kleur en vorm, maar ook naar opbouw, veerkracht en afwerking.

## Begin bij het frame

Het frame is de ruggengraat van de zetel. Je ziet het niet, maar het bepaalt mee hoe stabiel het meubel blijft na jaren gebruik. Vraag daarom welk materiaal gebruikt wordt en hoe het frame geconstrueerd is. Een degelijk houten frame of een geloofwaardige combinatie van stevige materialen geeft meer vertrouwen dan een vaag antwoord over 'sterke constructie'.

Als een verkoper nauwelijks kan uitleggen waaruit het frame bestaat, is dat op zich al informatie.

## Let op de vering

Vering bepaalt mee hoe een zetel terugveert en hoe lang hij zijn comfort behoudt. Een zetel die meteen wegzakt, voelt misschien heerlijk in minuut één, maar hoeft dat niet te blijven. Kijk daarom naar hoe de zitting reageert wanneer je je gewicht verplaatst of opstaat.

Goede vering voelt gecontroleerd. Niet plankhard, maar ook niet week. Je merkt dat de zetel steun geeft en terugkomt zonder futloos te worden.

## Vulling zegt veel over levensduur

Ook de samenstelling van de zit- en rugkussens maakt een groot verschil. Vraag naar de opbouw: schuim, dons, vezel of een combinatie. Niet om technische details te verzamelen, maar om te begrijpen hoe de zetel zal verouderen.

Een zetel met mooie eerste zachtheid maar weinig veerkracht kan snel slordig ogen. Een iets stevigere vulling blijft vaak beter in vorm. Dat wil niet zeggen dat stevig altijd beter is, wel dat je moet weten welk comforttype je koopt.

## Naden, afwerking en symmetrie

Kijk van dichtbij. Zijn naden netjes? Lopen lijnen mooi door? Sluiten kussens goed aan? Zitten er geen vreemde plooien of spanningen op plekken waar ze niet horen? Juist die details vertellen vaak meer dan de algemene indruk.

Een kwaliteitszetel voelt meestal coherent. Alles lijkt bewust: de vorm, de afwerking, de spanning van de stof, de balans tussen comfort en netheid.

> Kwaliteit verraadt zich zelden in één spectaculair detail. Je voelt ze juist in hoe weinig er wringt.

## Stel de juiste vragen

Goede verkopers herkennen goede vragen. Vraag dus niet alleen naar promoties of levertermijnen, maar ook naar opbouw, onderhoud, garantie en wat je realistisch mag verwachten qua gebruikssporen. Dat maakt het gesprek meteen inhoudelijker.

Zinnige vragen zijn bijvoorbeeld: welk type vering zit hierin? Hoe is de zitting opgebouwd? Hoe onderhoud ik deze bekleding? Wat gebeurt er visueel met deze zetel na intensief gebruik? Zulke vragen filteren showroomverkoop van echt advies.

## Vertrouw niet blind op prijs of merk

Een duurdere zetel kan degelijker zijn, maar prijs is geen sluitend bewijs. Ook een bekende naam garandeert niet automatisch dat elk model even goed opgebouwd is. Kijk dus altijd naar het concrete meubel voor je.

Omgekeerd kan een minder uitgesproken model soms verrassend degelijk zijn. Daarom blijft testen, kijken en vragen stellen belangrijker dan vertrouwen op etiket of uitstraling alleen.

## Wat doe je best in de winkel?

Ga zitten zoals thuis. Voel hoe de zetel reageert. Kijk hoe snel kussens herstellen. Duw licht op armleuningen en rug. Controleer hoe stabiel alles voelt. Bekijk de achterkant en zijkanten, niet alleen de ideale showroomhoek. Een goede zetel moet in het geheel overtuigen, niet alleen frontaal.

## Conclusie

Een kwalitatieve zetel herken je aan een combinatie van stabiele opbouw, geloofwaardig comfort, verzorgde afwerking en transparante uitleg. Wie daar bewust op let, koopt minder op gevoel alleen en maakt veel minder kans op een dure ontgoocheling.

Dat maakt van kwaliteit geen abstract begrip, maar iets heel concreets: een zetel die niet alleen vandaag goed oogt, maar ook over jaren nog goed zit.
""".strip(),
    ),
    (
        "Afmetingenartikel",
        """
# Artikel

**Titel:** Welke afmeting van zetel werkt in jouw woonkamer?  
**Slug:** /advies/zetel-afmetingen-woonkamer  
**Primaire zoekintentie:** praktische afmetingen en ruimteplanning  
**Primaire keywordfocus:** zetel afmetingen woonkamer, hoe groot mag een zetel zijn, hoekzetel kleine woonkamer  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Welke zetel past bij jouw leven?](https://interieur.expert/advies/welke-zetel-past-bij-jou), [Een zetel testen in de winkel](https://interieur.expert/advies/zetel-testen-in-de-winkel)

**Excerpt:** De juiste zetelmaat hangt niet alleen af van je woonkamer, maar ook van doorloop, zichtlijnen en gebruik. Zo bepaal je welke afmeting echt werkt in jouw ruimte.

Een zetel kan op zich perfect zijn en toch fout vallen in je woonkamer. Niet omdat hij lelijk is, maar omdat de maat niet klopt met de ruimte errond. Dat is een van de meest voorkomende oorzaken van spijt na aankoop. Mensen meten de muur, kiezen een model dat ongeveer past en vergeten alles daartussen: loopruimte, salontafelafstand, zichtlijnen en hoe zwaar een meubel optisch doorweegt.

## Breedte is maar één deel van het verhaal

De meeste mensen kijken eerst naar breedte. Dat is logisch, maar onvolledig. Een zetel met brede armleuningen of een diepe zitting kan veel massiever uitvallen dan de cijfers suggereren. Ook hoogte speelt mee. Een lage zetel kan visueel rustig zijn, maar in een kleine ruimte toch veel vloervlak blokkeren als hij diep is.

## Denk in gebruiksruimte, niet alleen in meubelmaat

Een zetel moet niet alleen passen, je moet er ook rond kunnen leven. Kun je nog gemakkelijk naar het raam? Blijft de doorgang naar de eetruimte logisch? Is er genoeg ruimte tussen zetel en salontafel? Kan iemand passeren zonder dat alles gepropt aanvoelt?

Zodra je in gebruiksruimte denkt, vermijd je veel klassieke fouten. Je kiest dan niet de grootste zetel die nog net binnen kan, maar de grootste zetel die de kamer nog goed laat functioneren.

## Wat werkt in een kleine woonkamer?

In kleine woonkamers winnen modellen met fijne poten, compacte armleuningen en heldere lijnen vaak. Ze nemen niet per se veel minder plaats in, maar ogen wel luchtiger. Dat verschil is belangrijk. Een zetel kan fysiek passen en visueel toch te zwaar aanvoelen.

Vaak werkt een compacte 3-zit beter dan een hoekzetel die de ruimte afsluit. Zeker wanneer je extra zitplaatsen met een fauteuil of lichte stoel kunt oplossen.

## Wanneer is een hoekzetel wél logisch?

Een hoekzetel werkt goed als je echt nood hebt aan veel zitplaatsen of wanneer de ruimte er structureel om vraagt. In open leefruimtes kan hij zones afbakenen en een sterke zithoek creëren. Maar in kleinere of druk ingedeelde kamers kan een hoekzetel snel dominant worden.

Vraag je dus af of je die extra hoek echt gebruikt, of vooral koopt omdat ze royaal oogt.

## De verhouding met andere meubels

Een zetel leeft nooit alleen. Hij moet samengaan met de salontafel, het tapijt, de verlichting en de rest van de zitruimte. Een te grote zetel maakt alles errond klein. Een te kleine zetel kan het omgekeerde doen: de kamer mist dan anker en samenhang.

Daarom is de beste maat meestal de maat die zowel zitcomfort als visueel evenwicht brengt.

> Een goede zetel vult een woonkamer niet op. Hij organiseert haar.

## Praktische vuistregels

Houd voldoende loopruimte rond de zetel zodat beweging natuurlijk blijft. Kies liever voor logische proporties dan voor maximale zitcapaciteit. En teken indien mogelijk de opstelling uit op de vloer met tape of papier. Dat klinkt basic, maar voorkomt verrassend veel miskopen.

## Conclusie

De juiste zetelmaat is niet de grootste maat die past, maar de maat die je woonkamer beter laat werken. Kijk dus niet alleen naar centimeters, maar ook naar doorstroming, visueel gewicht en dagelijkse logica. Dat is uiteindelijk wat het verschil maakt tussen een zetel die binnen geraakt en een zetel die echt klopt.
""".strip(),
    ),
    (
        "Showroomtestartikel",
        """
# Artikel

**Titel:** Een zetel testen in de winkel: 9 dingen die je echt moet doen  
**Slug:** /advies/zetel-testen-in-de-winkel  
**Primaire zoekintentie:** praktische koopchecklist  
**Primaire keywordfocus:** zetel testen in winkel, waar op letten bij zetel testen, showroom zetel tips  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Hoe herken je een kwalitatieve zetel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen), [Video: tips voor het kopen van een zetel](https://interieur.expert/video/tips-voor-het-kopen-van-een-zetel)

**Excerpt:** In de winkel even gaan zitten is niet genoeg. Met deze 9 praktische tests voel je sneller of een zetel thuis echt goed zal zitten en zijn prijs waard is.

Een showroom is gemaakt om een zetel aantrekkelijk te laten lijken. Het licht zit goed, de styling klopt en je ziet het meubel in ideale omstandigheden. Daardoor voelen veel modellen beter aan dan ze in dagelijks gebruik zullen doen. Wie slim wil kopen, test dus niet beleefd, maar bewust.

## 1. Ga zitten zoals je thuis zit

Zit niet alleen netjes rechtop op de rand. Doe wat je thuis ook doet. Leun achterover, schuif wat onderuit, verander van houding. Zo voel je veel sneller of de zetel natuurlijk zit of alleen in een ideale houding werkt.

## 2. Let op je voeten en knieën

Wanneer je zit, moeten je voeten logisch op de grond komen en je knieën niet ongemakkelijk omhoog of te laag uitkomen. Dat zegt veel over zithoogte en algemeen comfort.

## 3. Test de rugsteun bewust

Een goede rug voelt niet dwingend, maar wel ondersteunend. Merk je dat je meteen moet corrigeren met een kussen of jezelf actief moet recht houden, dan klopt de rugvorm wellicht niet voor jou.

## 4. Controleer de zitdiepte

Een diepe zetel kan heerlijk zijn als je graag loungt, maar lastiger wanneer je vaak rechtop zit. Voel dus of de diepte klopt met je lichaamsbouw en gebruik.

## 5. Sta meerdere keren op

Een zetel kan comfortabel lijken zolang je erin zit, maar onpraktisch blijken zodra je opstaat. Zeker bij erg lage of zachte modellen merk je dat snel. Test dat dus bewust.

## 6. Voel hoe de kussens herstellen

Komt de zitting redelijk terug in vorm nadat je opstaat? Of blijft ze meteen slordig ogen? Dat zegt veel over vulling en duurzaamheid.

## 7. Beoordeel armleuningen en randzones

Gebruik je de armleuningen echt? Leg er je arm op, steun erop en kijk of de hoogte logisch is. Ook de buitenste zitplaatsen verdienen aandacht, want net daar merk je vaak of een model doordacht is.

## 8. Kijk niet alleen naar de voorkant

Loop rond de zetel. Bekijk de achterkant, zijkanten en afwerking. Een goede zetel moet overal overtuigen, niet alleen vanuit de showroomhoek.

## 9. Vergelijk direct met een tweede model

Comfort beoordeel je beter in vergelijking. Test twee of drie modellen vlak na elkaar. Dan voel je veel sneller welke zetel echt beter zit en welke vooral een goede eerste indruk maakt.

> De beste showroomtest duurt langer dan twee minuten en voelt een beetje onelegant. Precies daarom is ze nuttig.

## Combineer dit met de video

Wie de showroomtest wil vertalen naar concrete koopvragen, kan hier perfect de videopagina naast zetten: [Video: tips voor het kopen van een zetel](https://interieur.expert/video/tips-voor-het-kopen-van-een-zetel). Die werkt als menselijk en geloofwaardig verlengstuk van dit artikel.

## Conclusie

Een zetel testen is meer dan even voelen of hij zacht zit. Het gaat om houding, steun, opstaan, afwerking en vergelijking. Wie dat bewust doet, maakt veel meer kans op een aankoop die thuis ook echt goed blijkt.
""".strip(),
    ),
    (
        "Videopagina",
        """
# Video

**Titel:** Video: tips voor het kopen van een zetel volgens een verkoper  
**Slug:** /video/tips-voor-het-kopen-van-een-zetel  
**YouTube-bron:** https://www.youtube.com/watch?v=rFmfG0PID9c  
**Primaire zoekintentie:** praktische video-uitleg + commerciële informatie  
**Primaire keywordfocus:** tips zetel kopen, waar op letten bij zetel kopen video, zetel kopen advies  
**Interne links:** [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten), [Een zetel testen in de winkel](https://interieur.expert/advies/zetel-testen-in-de-winkel), [Hoe herken je een kwalitatieve zetel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen)

**Excerpt:** In deze video deelt een verkoper uit de meubelretail praktische tips over waar je op moet letten wanneer je een nieuwe zetel kiest, test en vergelijkt.

Wie een zetel koopt, krijgt vaak veel showroominformatie tegelijk: stoffen, modellen, promoties, afmetingen, comfortopties. Net daarom werkt video hier goed. Je hoort en ziet advies in een concretere context, en dat maakt de aankoop voor veel mensen tastbaarder.

In deze video deelt een verkoper uit de meubelwereld praktische inzichten die veel kopers pas laat ontdekken. Niet als harde verkooppraat, maar als uitleg over wat in de winkel echt opvalt wanneer mensen een verkeerde of net een heel goede keuze maken.

## Wat je uit deze video haalt

De video helpt vooral op drie vlakken. Eerst krijg je extra gevoel bij hoe je een zetel moet testen in plaats van alleen snel uitproberen. Daarna leer je beter inschatten welke vragen in een winkel echt relevant zijn. En tenslotte helpt de video om theorie uit onze artikels sneller te vertalen naar een concreet koopmoment.

## Waarom deze video zinvol is binnen het cluster

Mensen die op dit punt in hun zoektocht zitten, hebben vaak al een eerste shortlist. Ze zoeken geen algemene inspiratie meer, maar praktische bevestiging. Video verlaagt dan de drempel. Het maakt advies menselijker, geloofwaardiger en makkelijker verteerbaar.

> Een goede video vervangt geen grondig artikel, maar haalt wel veel twijfel uit de laatste fase voor aankoop.

## Bekijk ook

- [Zetel kopen: waar moet je echt op letten?](https://interieur.expert/advies/zetel-kopen-waarop-letten)
- [Een zetel testen in de winkel: 9 dingen die je echt moet doen](https://interieur.expert/advies/zetel-testen-in-de-winkel)
- [Hoe herken je een kwalitatieve zetel in de winkel?](https://interieur.expert/advies/kwalitatieve-zetel-herkennen)

## Suggested supporting intro for the published video page

Een zetel koop je niet elke maand. Net daarom loont het om vóór je beslist goed te kijken naar comfort, opbouw, formaat en gebruik. In deze video deelt een verkoper uit de meubelretail praktische kooptips die helpen om bewuster te vergelijken en minder snel een miskoop te doen.
""".strip(),
    ),
]


def add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_inline_content(paragraph: Paragraph, text: str) -> None:
    position = 0
    for match in INLINE_LINK_RE.finditer(text):
        before = text[position:match.start()]
        add_bold_runs(paragraph, before)
        add_hyperlink(paragraph, match.group(1), match.group(2))
        position = match.end()
    add_bold_runs(paragraph, text[position:])


def add_bold_runs(paragraph: Paragraph, text: str) -> None:
    position = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])



def write_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_content(paragraph, line[2:].strip())
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_content(paragraph, re.sub(r"^\d+\. ", "", line))
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Intense Quote")
            add_inline_content(paragraph, line[2:].strip())
            continue
        paragraph = document.add_paragraph()
        add_inline_content(paragraph, line)



def set_base_styles(document: Document) -> None:
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
    document.styles["Heading 1"].font.size = Pt(20)
    document.styles["Heading 2"].font.size = Pt(15)
    document.styles["Heading 3"].font.size = Pt(12)



def add_page_break(document: Document) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)



def build_document() -> None:
    document = Document()
    set_base_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    document.add_heading(DOC_TITLE, level=0)
    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Interieur Expert - researchgedreven cluster rond zetels kopen in Vlaanderen")
    subtitle_run.italic = True

    write_markdown(document, INTRO)

    for index, (label, markdown) in enumerate(SECTIONS, start=1):
        add_page_break(document)
        label_paragraph = document.add_paragraph()
        label_run = label_paragraph.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)
        write_markdown(document, markdown)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(f"Created {OUTPUT_PATH}")
